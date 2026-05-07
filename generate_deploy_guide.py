"""
教师助手系统 - 云服务器部署教程 Word 文档生成脚本
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# ── 辅助函数 ──
def add_code_block(text, language=""):
    """添加带底色的代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 灰色底色
    shading = p.paragraph_format.element.get_or_add_pPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_tip(text, icon="💡"):
    """添加提示/注意框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{icon} {text}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.italic = True
    return p

def add_step(num, title):
    """添加步骤标题"""
    p = doc.add_paragraph()
    run = p.add_run(f"步骤 {num}：{title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0B, 0x66, 0x23)
    return p

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("教师助手系统\n云服务器部署教程")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("从零开始，手把手教你把项目部署到云服务器")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f"适用系统：Ubuntu 20.04 / 22.04 / 24.04 LTS\n生成日期：{datetime.date.today().strftime('%Y年%m月%d日')}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════════════════════════
doc.add_heading("目  录", level=1)
toc_items = [
    "一、项目架构总览",
    "二、购买与准备云服务器",
    "三、连接服务器（SSH 远程登录）",
    "四、安装基础环境（Java 17 + MySQL 8 + Node.js + Nginx）",
    "五、上传项目文件到服务器",
    "六、初始化数据库",
    "七、后端打包与启动",
    "八、前端构建与 Nginx 部署",
    "九、配置域名与 HTTPS（可选但推荐）",
    "十、设置开机自启动（让系统更稳定）",
    "十一、常见问题排查（FAQ）",
    "十二、部署后验证清单",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 一、项目架构总览
# ═══════════════════════════════════════════════════════════════
doc.add_heading("一、项目架构总览", level=1)

doc.add_paragraph(
    "教师助手系统是一个前后端分离的 Web 应用，部署到云服务器后，用户只需通过浏览器即可访问使用。"
    "在开始部署之前，请先了解项目的整体架构："
)

add_table(
    ["组件", "技术栈", "作用", "默认端口"],
    [
        ["前端", "Vue 3 + Element Plus + Vite", "用户界面，浏览器访问", "80（Nginx 代理）"],
        ["后端", "Spring Boot 3.5 + Java 17", "业务逻辑、AI 接口调用", "8089"],
        ["数据库", "MySQL 8.x", "存储用户、配置、历史等数据", "3306"],
        ["反向代理", "Nginx", "统一入口，静态文件 + API 转发", "80"],
    ]
)

doc.add_paragraph()
doc.add_paragraph("部署后的架构示意：")

add_code_block(
    "用户浏览器\n"
    "    │\n"
    "    ▼\n"
    "  Nginx（80 端口）\n"
    "    │\n"
    "    ├── 静态文件（前端 HTML/JS/CSS）──► /www/teacher-agent/frontend/dist/\n"
    "    │\n"
    "    └── /api/* 请求 ──► 反向代理到后端 Spring Boot（8089 端口）\n"
    "                          │\n"
    "                          └── 连接 MySQL（3306 端口）"
)

doc.add_paragraph()
doc.add_paragraph("你需要准备的东西：")
items = [
    "一台云服务器（推荐 2 核 4G 内存及以上配置）",
    "一个域名（可选，也可以直接用服务器 IP 访问）",
    "SSH 远程连接工具（Windows 自带的终端、PuTTY、或 Xshell 等）",
    "本项目的完整代码文件",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 二、购买与准备云服务器
# ═══════════════════════════════════════════════════════════════
doc.add_heading("二、购买与准备云服务器", level=1)

doc.add_heading("2.1 推荐云服务器配置", level=2)
add_table(
    ["配置项", "最低要求", "推荐配置", "说明"],
    [
        ["CPU", "2 核", "2 核", "AI 接口调用需要一定计算能力"],
        ["内存", "2 GB", "4 GB", "Java 程序至少需要 512MB，推荐留余量"],
        ["硬盘", "40 GB", "50 GB SSD", "用于存储系统、数据库和上传文件"],
        ["带宽", "1 Mbps", "3 Mbps 以上", "用户多的话需要更高带宽"],
        ["操作系统", "Ubuntu 20.04/22.04 LTS", "Ubuntu 22.04 LTS", "本教程以 Ubuntu 为例"],
    ]
)

doc.add_paragraph()
doc.add_paragraph("国内常见的云服务商：阿里云、腾讯云、华为云、百度云等，购买流程基本相同。")

doc.add_heading("2.2 购买后在控制台做的事", level=2)
steps = [
    "找到你的服务器公网 IP（类似 123.45.67.89），记下来，后面要用。",
    "设置服务器登录密码（在控制台「重置密码」功能中设置）。",
    "配置安全组（防火墙规则），开放以下端口：",
]
for s in steps:
    doc.add_paragraph(s)

add_table(
    ["端口", "协议", "用途", "是否必须"],
    [
        ["22", "TCP", "SSH 远程登录", "是"],
        ["80", "TCP", "HTTP 网站访问（Nginx）", "是"],
        ["443", "TCP", "HTTPS 安全访问（配置 SSL 后）", "建议"],
        ["3306", "TCP", "MySQL 数据库", "不需要对外开放，仅本机访问"],
        ["8089", "TCP", "后端 Spring Boot", "不需要对外开放，由 Nginx 转发"],
    ]
)

add_tip("安全组只需开放 22（SSH）和 80/443（HTTP/HTTPS）端口。后端 8089 和 MySQL 3306 不需要对公网开放，这样更安全。Nginx 在服务器内部会把 /api 请求转发到 8089 端口。", "🔒")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 三、连接服务器
# ═══════════════════════════════════════════════════════════════
doc.add_heading("三、连接服务器（SSH 远程登录）", level=1)

doc.add_paragraph("你需要用 SSH 工具连接到服务器。以下是几种常见方式：")

doc.add_heading("3.1 Windows 自带终端（推荐，最简单）", level=2)
doc.add_paragraph("Windows 10/11 自带 SSH 客户端，直接打开「终端」或「PowerShell」即可：")
add_code_block("ssh root@你的服务器IP")
add_tip("把「你的服务器IP」替换成实际的公网 IP，例如 ssh root@123.45.67.89")
doc.add_paragraph("第一次连接会提示是否信任服务器指纹，输入 yes 回车，然后输入密码（密码输入时屏幕不会显示任何字符，这是正常的），回车即可登录。")

doc.add_heading("3.2 使用 PuTTY / Xshell / Tabby 等工具", level=2)
doc.add_paragraph("如果你更喜欢图形界面工具，可以下载 PuTTY（免费）或 Xshell：")
items = [
    "PuTTY：下载后打开，Host Name 填服务器 IP，Port 填 22，点 Open",
    "Xshell：新建会话，主机填服务器 IP，端口 22，用户名 root",
    "Tabby：跨平台终端，界面美观，也支持 SSH 连接",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_tip("后续教程中所有以「$」开头的命令行，都是在 SSH 连接到服务器后的终端中执行的。复制命令时不要复制 $ 符号本身。", "📌")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 四、安装基础环境
# ═══════════════════════════════════════════════════════════════
doc.add_heading("四、安装基础环境", level=1)
doc.add_paragraph("连接到服务器后，依次安装以下软件。每个步骤都有完整的命令，直接复制粘贴执行即可。")

# ── 4.1 系统更新 ──
doc.add_heading("4.1 更新系统软件包", level=2)
doc.add_paragraph("安装任何软件之前，先更新系统到最新状态：")
add_code_block(
    "$ sudo apt update && sudo apt upgrade -y"
)
add_tip("这个命令会下载更新列表并升级所有已安装的软件包，可能需要几分钟。")

# ── 4.2 安装 JDK 17 ──
doc.add_heading("4.2 安装 JDK 17（Java 运行环境）", level=2)
doc.add_paragraph("后端是 Java 17 编写的 Spring Boot 项目，需要安装 JDK 17：")
add_code_block(
    "$ sudo apt install openjdk-17-jdk -y\n"
    "$ java -version"
)
doc.add_paragraph("如果安装成功，会看到类似输出：")
add_code_block("openjdk version \"17.0.x\" 2024-xx-xx")

# ── 4.3 安装 MySQL 8 ──
doc.add_heading("4.3 安装 MySQL 8（数据库）", level=2)
doc.add_paragraph("安装 MySQL 数据库服务器：")
add_code_block(
    "$ sudo apt install mysql-server -y\n"
    "$ sudo systemctl start mysql\n"
    "$ sudo systemctl enable mysql"
)
doc.add_paragraph("验证 MySQL 是否运行：")
add_code_block("$ sudo systemctl status mysql")
doc.add_paragraph("看到「active (running)」就说明 MySQL 已经正常运行。")

doc.add_paragraph()
doc.add_heading("配置 MySQL root 密码和安全设置：", level=3)
add_code_block(
    "$ sudo mysql\n"
    "\n"
    "mysql> ALTER USER 'root'@'localhost' IDENTIFIED WITH mysql_native_password BY '你的数据库密码';\n"
    "mysql> FLUSH PRIVILEGES;\n"
    "mysql> EXIT;"
)
add_tip("请把「你的数据库密码」改成一个你记得住的强密码，后面配置后端时要用到。建议包含大小写字母和数字，例如：TeacherAgent2026!", "🔐")

# ── 4.4 安装 Node.js ──
doc.add_heading("4.4 安装 Node.js（前端构建工具）", level=2)
doc.add_paragraph("Node.js 只在构建前端时使用，部署后不需要一直运行。")
add_code_block(
    "# 安装 Node.js 20.x LTS\n"
    "$ curl -fsSL https://deb.nodesource.com/setup_20.x | sudo -E bash -\n"
    "$ sudo apt install -y nodejs\n"
    "\n"
    "# 验证安装\n"
    "$ node -v\n"
    "$ npm -v"
)
add_tip("如果 curl 命令不存在，先执行 sudo apt install curl -y 安装它。")

# ── 4.5 安装 Nginx ──
doc.add_heading("4.5 安装 Nginx（Web 服务器 / 反向代理）", level=2)
doc.add_paragraph("Nginx 负责给用户提供前端页面，同时把 API 请求转发给后端：")
add_code_block(
    "$ sudo apt install nginx -y\n"
    "$ sudo systemctl start nginx\n"
    "$ sudo systemctl enable nginx"
)
doc.add_paragraph("此时在浏览器中访问 http://你的服务器IP，应该能看到 Nginx 的默认欢迎页面。")

# ── 4.6 安装 Maven ──
doc.add_heading("4.6 安装 Maven（Java 构建工具）", level=2)
doc.add_paragraph("Maven 用于打包后端 Java 项目：")
add_code_block(
    "$ sudo apt install maven -y\n"
    "$ mvn -version"
)
add_tip("如果 apt 安装的 Maven 版本过低，也可以手动安装：下载 Apache Maven 3.9.x，解压到 /opt/maven，然后配置 PATH 环境变量。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 五、上传项目文件
# ═══════════════════════════════════════════════════════════════
doc.add_heading("五、上传项目文件到服务器", level=1)

doc.add_heading("5.1 创建项目目录", level=2)
add_code_block(
    "$ sudo mkdir -p /www/teacher-agent\n"
    "$ cd /www/teacher-agent"
)

doc.add_heading("5.2 上传方式一：从本地 Windows 上传（推荐新手）", level=2)
doc.add_paragraph("使用 SCP 命令上传（在本地 Windows 终端执行，不是在服务器上执行）：")

add_code_block(
    "# 在本地 Windows PowerShell 或 CMD 中执行\n"
    "# 上传整个项目文件夹\n"
    "scp -r D:\\AI-project\\teacher-agent\\backend root@你的服务器IP:/www/teacher-agent/\n"
    "scp -r D:\\AI-project\\teacher-agent\\frontend root@你的服务器IP:/www/teacher-agent/"
)

add_tip("上传可能需要几分钟，取决于网速。如果你用 PuTTY，可以配合 PSFTP 或 WinSCP 工具上传文件，有图形界面更直观。", "📁")

doc.add_heading("5.3 上传方式二：使用 Git（如果代码在 GitHub/Gitee 上）", level=2)
add_code_block(
    "$ cd /www/teacher-agent\n"
    "$ sudo apt install git -y\n"
    "$ git clone 你的仓库地址 ."
)

doc.add_heading("5.4 上传完成后的目录结构", level=2)
add_code_block(
    "/www/teacher-agent/\n"
    "├── backend/              # 后端 Spring Boot 项目\n"
    "│   ├── pom.xml\n"
    "│   ├── mvnw\n"
    "│   └── src/\n"
    "├── frontend/             # 前端 Vue 项目\n"
    "│   ├── package.json\n"
    "│   └── src/\n"
    "└── data/                 # 运行时数据目录（自动创建）"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 六、初始化数据库
# ═══════════════════════════════════════════════════════════════
doc.add_heading("六、初始化数据库", level=1)

doc.add_heading("6.1 创建数据库", level=2)
doc.add_paragraph("登录 MySQL 并创建项目所需的数据库：")
add_code_block(
    "$ sudo mysql -u root -p\n"
    "\n"
    "# 输入你之前设置的 root 密码，然后执行以下 SQL：\n"
    "\n"
    "CREATE DATABASE teacher_agent DEFAULT CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;\n"
    "\n"
    "# 创建一个专用数据库用户（更安全，不用 root 直接连接）\n"
    "CREATE USER 'teacher_agent'@'localhost' IDENTIFIED BY '你的数据库密码';\n"
    "GRANT ALL PRIVILEGES ON teacher_agent.* TO 'teacher_agent'@'localhost';\n"
    "FLUSH PRIVILEGES;\n"
    "EXIT;"
)

doc.add_heading("6.2 导入数据表结构", level=2)
doc.add_paragraph("项目的 SQL 迁移脚本在 backend/src/main/resources/db/migration/ 目录下。依次执行：")
add_code_block(
    "$ cd /www/teacher-agent\n"
    "\n"
    "# 导入 V2 表结构\n"
    "$ sudo mysql -u teacher_agent -p teacher_agent < backend/src/main/resources/db/migration/V2__course_config.sql\n"
    "\n"
    "# 导入 V3 表结构\n"
    "$ sudo mysql -u teacher_agent -p teacher_agent < backend/src/main/resources/db/migration/V3__knowledge_base.sql"
)
add_tip("如果还有其他 SQL 文件（如 V1_xxx.sql），也要按版本号顺序导入。每次导入都会提示输入密码。")

doc.add_heading("6.3 确认表是否创建成功", level=2)
add_code_block(
    "$ sudo mysql -u teacher_agent -p -e \"USE teacher_agent; SHOW TABLES;\""
)
doc.add_paragraph("你应该能看到 course_config、knowledge_base、knowledge_chunk 等表名。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 七、后端打包与启动
# ═══════════════════════════════════════════════════════════════
doc.add_heading("七、后端打包与启动", level=1)

doc.add_heading("7.1 配置数据库连接", level=2)
doc.add_paragraph("编辑后端的配置文件，把数据库密码改成你在第六步设置的密码：")
add_code_block(
    "$ sudo nano /www/teacher-agent/backend/src/main/resources/application.yml"
)

doc.add_paragraph("找到以下内容并修改：")
add_code_block(
    "spring:\n"
    "  datasource:\n"
    "    url: jdbc:mysql://localhost:3306/teacher_agent?useUnicode=true&characterEncoding=utf8&useSSL=false&serverTimezone=Asia/Shanghai&allowPublicKeyRetrieval=true&createDatabaseIfNotExist=true\n"
    "    username: ${TEACHER_AGENT_DB_USERNAME:root}\n"
    "    password: ${TEACHER_AGENT_DB_PASSWORD:你的数据库密码}"
)

doc.add_paragraph()
doc.add_paragraph("同时检查 JWT 密钥配置，建议改为随机字符串：")
add_code_block(
    "jwt:\n"
    "  secret: ${JWT_SECRET:改成你自己的随机密钥至少32个字符}\n"
    "  expire-hours: 168"
)

add_tip("nano 编辑器使用方法：用方向键移动光标 → 修改内容 → Ctrl+O 保存 → 回车确认 → Ctrl+X 退出。", "📝")

doc.add_heading("7.2 打包后端项目", level=2)
doc.add_paragraph("使用 Maven 将 Java 项目打包成可执行的 JAR 文件：")
add_code_block(
    "$ cd /www/teacher-agent/backend\n"
    "$ mvn clean package -DskipTests"
)
add_tip("首次打包会下载大量依赖包（约 200MB），需要等待 5-15 分钟。如果网络慢，可以配置阿里云 Maven 镜像加速。打包成功后会在 backend/target/ 目录生成 teacher-agent.jar 文件。")

doc.add_heading("7.3 启动后端服务", level=2)
doc.add_paragraph("先创建必要的目录：")
add_code_block(
    "$ cd /www/teacher-agent/backend\n"
    "$ mkdir -p data/uploads data/outputs data/templates data/materials"
)

doc.add_paragraph("启动后端（前台启动，用于测试）：")
add_code_block(
    "$ cd /www/teacher-agent/backend\n"
    "$ java -Xmx512m -Dfile.encoding=UTF-8 -Dnative.encoding=UTF-8 -jar target/teacher-agent.jar"
)

doc.add_paragraph("看到类似以下日志说明启动成功：")
add_code_block(
    "Started TeacherAgentApplication in x.xxx seconds"
)

add_tip("如果一切正常，按 Ctrl+C 先停止，接下来用后台方式运行。")

doc.add_heading("7.4 后台运行（长期运行方式）", level=2)
doc.add_paragraph("使用 nohup 让后端在后台运行，即使关闭终端也不会停止：")
add_code_block(
    "$ cd /www/teacher-agent/backend\n"
    "$ nohup java -Xmx512m -Dfile.encoding=UTF-8 -Dnative.encoding=UTF-8 \\\n"
    "  -jar target/teacher-agent.jar \\\n"
    "  > logs/app.log 2>&1 &\n"
    "\n"
    "# 先创建日志目录\n"
    "$ mkdir -p logs"
)

doc.add_paragraph("重新整理完整命令：")
add_code_block(
    "$ cd /www/teacher-agent/backend\n"
    "$ mkdir -p logs\n"
    "$ nohup java -Xmx512m -Dfile.encoding=UTF-8 -Dnative.encoding=UTF-8 -jar target/teacher-agent.jar > logs/app.log 2>&1 &"
)

doc.add_paragraph("查看后端是否正在运行：")
add_code_block(
    "$ tail -f logs/app.log          # 查看实时日志（Ctrl+C 退出查看）\n"
    "$ curl http://localhost:8089     # 测试后端是否响应"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 八、前端构建与 Nginx 部署
# ═══════════════════════════════════════════════════════════════
doc.add_heading("八、前端构建与 Nginx 部署", level=1)

doc.add_heading("8.1 安装前端依赖并构建", level=2)
add_code_block(
    "$ cd /www/teacher-agent/frontend\n"
    "$ npm install\n"
    "$ npm run build"
)
add_tip("构建过程需要几分钟。成功后会在 frontend/dist/ 目录生成静态文件（HTML、JS、CSS）。")

doc.add_heading("8.2 配置 Nginx", level=2)
doc.add_paragraph("创建 Nginx 站点配置文件：")
add_code_block(
    "$ sudo nano /etc/nginx/sites-available/teacher-agent"
)

doc.add_paragraph("写入以下完整配置（请根据实际情况修改）：")
add_code_block(
    "server {\n"
    "    listen 80;\n"
    "    server_name 你的域名或服务器IP;\n"
    "\n"
    "    # 前端静态文件\n"
    "    root /www/teacher-agent/frontend/dist;\n"
    "    index index.html;\n"
    "\n"
    "    # 前端路由（Vue Router 使用 hash 模式，这个 fallback 保证刷新不 404）\n"
    "    location / {\n"
    "        try_files $uri $uri/ /index.html;\n"
    "    }\n"
    "\n"
    "    # API 反向代理：把 /api 请求转发到后端 Spring Boot\n"
    "    location /api/ {\n"
    "        proxy_pass http://127.0.0.1:8089/api/;\n"
    "        proxy_set_header Host $host;\n"
    "        proxy_set_header X-Real-IP $remote_addr;\n"
    "        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n"
    "        proxy_set_header X-Forwarded-Proto $scheme;\n"
    "\n"
    "        # 文件上传大小限制（后端配置了 200MB）\n"
    "        client_max_body_size 200m;\n"
    "\n"
    "        # SSE 支持（任务进度实时推送需要）\n"
    "        proxy_http_version 1.1;\n"
    "        proxy_set_header Connection '';\n"
    "        proxy_buffering off;\n"
    "        proxy_cache off;\n"
    "        chunked_transfer_encoding on;\n"
    "        proxy_read_timeout 600s;\n"
    "    }\n"
    "\n"
    "    # 静态资源缓存（CSS/JS/图片等）\n"
    "    location ~* \\.(js|css|png|jpg|jpeg|gif|ico|svg|woff|woff2|ttf|eot)$ {\n"
    "        expires 30d;\n"
    "        add_header Cache-Control \"public, immutable\";\n"
    "    }\n"
    "}"
)

doc.add_paragraph("把配置链接到 Nginx 启用目录，并测试配置是否正确：")
add_code_block(
    "# 创建软链接\n"
    "$ sudo ln -s /etc/nginx/sites-available/teacher-agent /etc/nginx/sites-enabled/\n"
    "\n"
    "# 删除默认站点（可选，如果你只用这一个站点）\n"
    "$ sudo rm -f /etc/nginx/sites-enabled/default\n"
    "\n"
    "# 测试 Nginx 配置语法\n"
    "$ sudo nginx -t"
)

doc.add_paragraph("如果输出「syntax is ok / test is successful」，就说明配置没有问题。重载 Nginx：")
add_code_block("$ sudo systemctl reload nginx")

add_tip("如果在同一台服务器上还部署了其他网站，不要删除 default，而是用不同 server_name 区分。")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 九、域名与 HTTPS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("九、配置域名与 HTTPS（可选但推荐）", level=1)

doc.add_heading("9.1 绑定域名", level=2)
doc.add_paragraph(
    "如果你有自己的域名（例如 teacher.example.com），需要做两件事："
)
items = [
    "在域名服务商（如阿里云、腾讯云）的 DNS 解析中添加一条 A 记录，将域名指向你的服务器公网 IP。",
    "修改 Nginx 配置中的 server_name 为你的域名。",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("9.2 安装 SSL 证书（让网站更安全）", level=2)
doc.add_paragraph("使用 Let's Encrypt 免费 SSL 证书，一键配置 HTTPS：")
add_code_block(
    "# 安装 Certbot\n"
    "$ sudo apt install certbot python3-certbot-nginx -y\n"
    "\n"
    "# 自动获取并配置 SSL 证书\n"
    "$ sudo certbot --nginx -d 你的域名"
)
doc.add_paragraph("按提示输入邮箱地址（用于证书到期提醒），选择是否重定向 HTTP 到 HTTPS。")
add_tip("Let's Encrypt 证书有效期为 90 天，Certbot 会自动续期，无需手动操作。可以执行 sudo certbot renew --dry-run 测试续期流程。", "🔄")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 十、开机自启动
# ═══════════════════════════════════════════════════════════════
doc.add_heading("十、设置开机自启动（让系统更稳定）", level=1)

doc.add_paragraph(
    "如果服务器重启，我们希望后端服务能自动启动。使用 systemd 服务来实现。"
)

doc.add_heading("10.1 创建后端 systemd 服务", level=2)
add_code_block("$ sudo nano /etc/systemd/system/teacher-agent.service")

doc.add_paragraph("写入以下内容：")
add_code_block(
    "[Unit]\n"
    "Description=Teacher Agent Backend Service\n"
    "After=network.target mysql.service\n"
    "Requires=mysql.service\n"
    "\n"
    "[Service]\n"
    "Type=simple\n"
    "User=root\n"
    "WorkingDirectory=/www/teacher-agent/backend\n"
    "ExecStart=/usr/bin/java -Xmx512m -Dfile.encoding=UTF-8 -Dnative.encoding=UTF-8 -jar /www/teacher-agent/backend/target/teacher-agent.jar\n"
    "ExecStop=/bin/kill -15 $MAINPID\n"
    "Restart=on-failure\n"
    "RestartSec=10\n"
    "StandardOutput=append:/www/teacher-agent/backend/logs/app.log\n"
    "StandardError=append:/www/teacher-agent/backend/logs/error.log\n"
    "\n"
    "[Install]\n"
    "WantedBy=multi-user.target"
)

doc.add_heading("10.2 启用并启动服务", level=2)
add_code_block(
    "# 重载 systemd 配置\n"
    "$ sudo systemctl daemon-reload\n"
    "\n"
    "# 设置开机自启动\n"
    "$ sudo systemctl enable teacher-agent\n"
    "\n"
    "# 启动服务\n"
    "$ sudo systemctl start teacher-agent\n"
    "\n"
    "# 查看运行状态\n"
    "$ sudo systemctl status teacher-agent"
)

doc.add_paragraph("看到绿色的「active (running)」表示服务运行正常。")

doc.add_paragraph()
doc.add_paragraph("常用服务管理命令速查：")
add_table(
    ["操作", "命令"],
    [
        ["启动后端", "sudo systemctl start teacher-agent"],
        ["停止后端", "sudo systemctl stop teacher-agent"],
        ["重启后端", "sudo systemctl restart teacher-agent"],
        ["查看状态", "sudo systemctl status teacher-agent"],
        ["查看实时日志", "tail -f /www/teacher-agent/backend/logs/app.log"],
        ["查看错误日志", "tail -f /www/teacher-agent/backend/logs/error.log"],
    ]
)

add_tip("如果你之前用 nohup 方式启动了后端，需要先 kill 掉旧进程，再用 systemctl 方式启动。执行 ps aux | grep teacher-agent 找到进程 ID，然后 kill 进程ID。", "📌")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 十一、常见问题排查
# ═══════════════════════════════════════════════════════════════
doc.add_heading("十一、常见问题排查（FAQ）", level=1)

problems = [
    (
        "Q1：浏览器打开 IP 地址，看到「Welcome to nginx!」而不是项目页面",
        "说明 Nginx 还在用默认配置。检查：\n"
        "1. 确认 frontend/dist/ 目录下有 index.html 文件\n"
        "2. 确认 /etc/nginx/sites-enabled/teacher-agent 软链接存在\n"
        "3. 执行 sudo nginx -t 检查配置语法\n"
        "4. 执行 sudo systemctl reload nginx 重新加载"
    ),
    (
        "Q2：页面打开了但是点击功能报错「网络错误」",
        "后端可能没启动或者 Nginx 代理配置有误。排查步骤：\n"
        "1. 在服务器上执行 curl http://localhost:8089 看后端是否正常\n"
        "2. 如果后端正常，检查 Nginx 中 /api/ 的 proxy_pass 地址是否正确\n"
        "3. 查看后端日志：tail -100 /www/teacher-agent/backend/logs/app.log"
    ),
    (
        "Q3：后端启动报错「Communications link failure」或连不上数据库",
        "MySQL 连接问题。排查步骤：\n"
        "1. 确认 MySQL 正在运行：sudo systemctl status mysql\n"
        "2. 确认数据库存在：sudo mysql -e \"SHOW DATABASES;\"\n"
        "3. 检查 application.yml 中的用户名和密码是否正确\n"
        "4. 测试连接：mysql -u teacher_agent -p teacher_agent"
    ),
    (
        "Q4：上传文件时报错「文件太大」",
        "需要在 Nginx 配置中添加 client_max_body_size 设置（已在上面配置中包含 200m）。\n"
        "修改后执行 sudo systemctl reload nginx 重新加载。"
    ),
    (
        "Q5：AI 功能不工作（教案/题库生成失败）",
        "AI 功能需要配置 LLM（大语言模型）API 密钥。部署后需要：\n"
        "1. 用管理员账号登录系统\n"
        "2. 进入「系统配置」或「LLM 配置」页面\n"
        "3. 填入 DeepSeek / 智谱 / 通义千问等 API 密钥并保存"
    ),
    (
        "Q6：前端 npm run build 报错",
        "常见原因：\n"
        "1. Node.js 版本过低 —— 需要 18.x 以上，执行 node -v 检查\n"
        "2. 内存不足 —— 4G 以下服务器构建可能失败，可以临时增加交换空间\n"
        "3. 依赖下载失败 —— 可以配置 npm 淘宝镜像：npm config set registry https://registry.npmmirror.com"
    ),
    (
        "Q7：服务器重启后网站打不开了",
        "需要确认服务是否设置了开机自启动：\n"
        "1. 执行 sudo systemctl enable teacher-agent（后端）\n"
        "2. Nginx 和 MySQL 在安装时已设置 enable，通常无需额外操作\n"
        "3. 检查各服务状态：systemctl status nginx mysql teacher-agent"
    ),
]

for q, a in problems:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(a)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 十二、部署后验证清单
# ═══════════════════════════════════════════════════════════════
doc.add_heading("十二、部署后验证清单", level=1)

doc.add_paragraph("部署完成后，请逐项检查以下功能是否正常：")

checklist = [
    ("打开浏览器访问", "http://你的服务器IP 或 https://你的域名", "能看到登录页面"),
    ("注册新用户", "在登录页点击注册，填写信息", "注册成功，自动跳转首页"),
    ("管理员登录", "使用管理员账号登录", "能看到所有菜单"),
    ("配置 LLM", "进入「系统配置」→「LLM 配置」", "能填入 API 密钥并保存"),
    ("上传教材", "进入「知识库」→ 上传 PPT/PDF", "上传成功，能解析出知识点"),
    ("生成教案", "进入「教案生成」→ 选择章节 → 生成", "能成功生成并下载 Word 文档"),
    ("生成题库", "进入「题库生成」→ 选择章节 → 生成", "能成功生成并下载 Excel 文件"),
    ("AI 问答", "进入「AI 助教」→ 提问", "能收到 AI 回答（SSE 流式输出）"),
    ("历史记录", "查看教案/题库历史", "能看到历史生成记录"),
    ("文件下载", "点击下载按钮", "能正确下载 docx/xlsx 文件"),
]

add_table(
    ["序号", "检查项目", "操作方式", "预期结果"],
    [(str(i+1), *row) for i, row in enumerate(checklist)]
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("—— 部署教程结束 ——")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("如果在部署过程中遇到任何问题，请按照「常见问题排查」章节逐一检查。")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── 保存文档 ──
output_path = r"D:\AI-project\teacher-agent\教师助手系统_云服务器部署教程.docx"
doc.save(output_path)
import sys
sys.stdout.reconfigure(encoding='utf-8')
print(f"部署教程已生成: {output_path}")
